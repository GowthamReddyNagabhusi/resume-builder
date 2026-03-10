"""Dynamic resume assembler from stored profile data and user selections."""

from __future__ import annotations

from datetime import datetime
from pathlib import Path

from backend.database import models as db

try:
    from docx import Document
except Exception:
    Document = None

try:
    from reportlab.lib.pagesizes import A4
    from reportlab.pdfgen import canvas
except Exception:
    canvas = None

OUTPUT_DIR = Path(__file__).resolve().parents[2] / "resume"


def _join(values):
    return ", ".join([v for v in values if v])


def _build_content(user_id: int, config: dict) -> dict:
    bundle = db.get_profile_bundle(user_id)
    profile = bundle.get("profile") or {}

    projects = bundle.get("projects", [])
    if config.get("selected_projects"):
        project_ids = set(config["selected_projects"])
        projects = [p for p in projects if p.get("id") in project_ids]

    internships = bundle.get("internships", [])
    if config.get("selected_experience"):
        ex_ids = set(config["selected_experience"])
        internships = [e for e in internships if e.get("id") in ex_ids]

    coding_platforms = bundle.get("coding_platforms", [])
    if config.get("selected_platforms"):
        p_ids = set(config["selected_platforms"])
        coding_platforms = [p for p in coding_platforms if p.get("id") in p_ids]

    return {
        "profile": profile,
        "education": bundle.get("education", []),
        "projects": projects,
        "internships": internships,
        "certifications": bundle.get("certifications", []),
        "training": bundle.get("training", []),
        "coding_platforms": coding_platforms,
        "selected_skills": config.get("selected_skills", []),
        "target_role": config.get("target_role", ""),
    }


def _write_docx(content: dict, output_path: Path):
    if Document is None:
        raise RuntimeError("python-docx not installed")

    doc = Document()
    profile = content["profile"]

    doc.add_heading(profile.get("full_name") or "Developer Resume", 0)
    contact = _join([
        profile.get("email", ""),
        profile.get("phone", ""),
        profile.get("location", ""),
        profile.get("linkedin", ""),
        profile.get("portfolio", ""),
    ])
    if contact:
        doc.add_paragraph(contact)

    if content.get("target_role"):
        doc.add_heading("Target Role", level=1)
        doc.add_paragraph(content["target_role"])

    if content.get("selected_skills"):
        doc.add_heading("Skills", level=1)
        doc.add_paragraph(_join(content["selected_skills"]))

    doc.add_heading("Education", level=1)
    for e in content.get("education", []):
        doc.add_paragraph(
            f"{e.get('degree', '')} ({e.get('branch', '')}) - {e.get('university', '')} [{e.get('start_year', '')} - {e.get('end_year', '')}]"
        )

    doc.add_heading("Projects", level=1)
    for p in content.get("projects", []):
        para = doc.add_paragraph(style="List Bullet")
        para.add_run(p.get("title", "Untitled")).bold = True
        details = _join([
            p.get("description", ""),
            p.get("tech_stack", ""),
            p.get("github_link", ""),
            p.get("live_link", ""),
        ])
        if details:
            para.add_run(f": {details}")

    doc.add_heading("Experience", level=1)
    for x in content.get("internships", []):
        para = doc.add_paragraph(style="List Bullet")
        para.add_run(f"{x.get('role', '')} - {x.get('company', '')}").bold = True
        para.add_run(f" ({x.get('start_date', '')} to {x.get('end_date', '')})")
        if x.get("description"):
            doc.add_paragraph(x.get("description"))

    if content.get("certifications"):
        doc.add_heading("Certifications", level=1)
        for c in content["certifications"]:
            doc.add_paragraph(
                f"{c.get('certificate_name', '')} - {c.get('provider', '')} ({c.get('issue_date', '')})",
                style="List Bullet",
            )

    if content.get("coding_platforms"):
        doc.add_heading("Coding Platforms", level=1)
        for cp in content["coding_platforms"]:
            doc.add_paragraph(
                f"{cp.get('platform_name', '')}: {cp.get('username', '')} {cp.get('profile_link', '')}",
                style="List Bullet",
            )

    doc.save(str(output_path))


def _write_pdf(content: dict, output_path: Path):
    if canvas is None:
        raise RuntimeError("reportlab not installed")

    c = canvas.Canvas(str(output_path), pagesize=A4)
    y = 800

    def line(text: str, size: int = 10, bold: bool = False):
        nonlocal y
        font = "Helvetica-Bold" if bold else "Helvetica"
        c.setFont(font, size)
        c.drawString(40, y, text[:120])
        y -= 16

    profile = content["profile"]
    line(profile.get("full_name") or "Developer Resume", 16, True)
    line(_join([profile.get("email", ""), profile.get("phone", ""), profile.get("location", "")]), 10)
    line(" ")

    if content.get("target_role"):
        line("Target Role", 12, True)
        line(content["target_role"])

    line("Skills", 12, True)
    line(_join(content.get("selected_skills", [])))

    line("Projects", 12, True)
    for p in content.get("projects", []):
        line(f"- {p.get('title', '')}: {p.get('description', '')}")

    line("Experience", 12, True)
    for x in content.get("internships", []):
        line(f"- {x.get('role', '')} @ {x.get('company', '')}")

    c.save()


def build_dynamic_resume(user_id: int, config: dict, config_id: int) -> dict:
    content = _build_content(user_id, config)
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

    output_type = (config.get("output_type") or "docx").lower()
    target = (config.get("target_role") or "general").replace(" ", "_")
    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f"dynamic_u{user_id}_c{config_id}_{target}_{ts}.{output_type}"
    path = OUTPUT_DIR / filename

    if output_type == "pdf":
        _write_pdf(content, path)
    else:
        _write_docx(content, path)

    return {"file_path": str(path), "filename": filename}
