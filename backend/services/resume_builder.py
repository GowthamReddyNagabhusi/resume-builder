"""
backend/services/resume_builder.py
Smart AI-powered resume generation with intelligent project selection.
Uses Groq for bullet generation and python-docx for DOCX output.
"""

import json
from pathlib import Path
from datetime import datetime

from backend.services.ai_engine import generate_resume_bullets
from backend.database import models as db

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    DOCX_OK = True
except ImportError:
    DOCX_OK = False

NAME_COLOR    = RGBColor(0x2F, 0x54, 0x96)
SECTION_COLOR = RGBColor(0x1F, 0x38, 0x64)
LABEL_COLOR   = RGBColor(0x00, 0x33, 0xCC)
BORDER_HEX    = "2F5496"

OUTPUT_DIR = Path(__file__).parent.parent.parent / "resume"


# ── Smart project picker ────────────────────────────────────

def _select_best_projects(all_projects: list, max_count: int = 3) -> list:
    """
    Intelligently pick the best 2-3 GitHub projects for a resume.

    Scoring per project (100 pts max):
      - Stars:            up to 30 pts  (popularity signal)
      - Description:      up to 25 pts  (has actual content)
      - Recency:          up to 20 pts  (recently updated)
      - Real language:    up to 15 pts  (not null / N/A)
      - Name quality:     up to 10 pts  (not auto-named like "username.github.io")

    Language diversity: at most 2 projects sharing the same language.
    """
    if not all_projects:
        return []

    max_stars = max((p.get("stars", 0) for p in all_projects), default=1) or 1

    scored = []
    for proj in all_projects:
        score = 0.0

        # Stars (0-30)
        score += (proj.get("stars", 0) / max_stars) * 30

        # Description quality (0-25)
        desc = (proj.get("description") or "").strip()
        if len(desc) > 50:
            score += 25
        elif len(desc) > 20:
            score += 15
        elif len(desc) > 5:
            score += 5

        # Recency (0-20) — decays linearly over 365 days
        raw_date = proj.get("updated_at", "")
        if raw_date:
            try:
                updated  = datetime.fromisoformat(raw_date[:19])
                age_days = max(0, (datetime.now() - updated).days)
                score   += max(0.0, 20.0 - (age_days / 365) * 20.0)
            except Exception:
                pass

        # Has a real programming language (0-15)
        lang = (proj.get("language") or "").strip()
        if lang and lang.lower() not in ("", "n/a", "null"):
            score += 15

        # Name quality (0-10) — penalise profile repos / forks / boilerplates
        name = proj.get("name", "").lower()
        boring_patterns = (".github.io", "-fork", "todo", "hello-world", "test", "practice")
        if not any(p in name for p in boring_patterns):
            score += 10

        scored.append((score, proj))

    scored.sort(key=lambda x: x[0], reverse=True)

    # Pick top-N with language diversity (max 2 per language)
    selected   = []
    lang_count = {}
    for _, proj in scored:
        if len(selected) >= max_count:
            break
        lang = (proj.get("language") or "other").lower()
        if lang_count.get(lang, 0) < 2:
            selected.append(proj)
            lang_count[lang] = lang_count.get(lang, 0) + 1

    print(f"[Resume] Selected {len(selected)} best projects from {len(all_projects)} total")
    for p in selected:
        print(f"  - {p['name']} (stars={p.get('stars',0)}, lang={p.get('language','?')})")

    return selected


# ── DOCX helper functions ───────────────────────────────────

def _spacing(para, before=0, after=0):
    pPr = para._p.get_or_add_pPr()
    sp = OxmlElement("w:spacing")
    sp.set(qn("w:before"), str(before))
    sp.set(qn("w:after"), str(after))
    pPr.append(sp)


def _bottom_border(para):
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    b = OxmlElement("w:bottom")
    b.set(qn("w:val"), "single"); b.set(qn("w:sz"), "6")
    b.set(qn("w:space"), "1");   b.set(qn("w:color"), BORDER_HEX)
    pBdr.append(b); pPr.append(pBdr)


def _right_tab(para, inches=7.3):
    pPr = para._p.get_or_add_pPr()
    tabs = OxmlElement("w:tabs")
    tab  = OxmlElement("w:tab")
    tab.set(qn("w:val"), "right")
    tab.set(qn("w:pos"), str(int(inches * 1440)))
    tabs.append(tab); pPr.append(tabs)


def _section(doc, title):
    p = doc.add_paragraph()
    _spacing(p, before=100, after=40)
    _bottom_border(p)
    r = p.add_run(title)
    r.font.size = Pt(11)
    r.font.color.rgb = SECTION_COLOR
    return p


def _bullet(doc, text):
    if not text or not text.strip():
        return
    p = doc.add_paragraph(style="List Paragraph")
    _spacing(p, before=0, after=20)
    p.add_run(text.strip())


def _skill_line(doc, label, value):
    if not value:
        return
    p = doc.add_paragraph()
    _spacing(p, before=0, after=20)
    r1 = p.add_run(label)
    r1.bold = True
    r1.font.color.rgb = LABEL_COLOR
    p.add_run(": " + value)


def _project_header(doc, name, date):
    p = doc.add_paragraph()
    _spacing(p, before=80, after=20)
    _right_tab(p, 7.3)
    r1 = p.add_run("  " + name + "\t")
    r1.bold = True; r1.font.size = Pt(11); r1.font.color.rgb = LABEL_COLOR
    p.add_run(date).font.size = Pt(10)


def _edu_entry(doc, institution, location, degree, grade, period):
    p1 = doc.add_paragraph()
    _spacing(p1, before=40, after=0)
    _right_tab(p1, 7.3)
    r1 = p1.add_run(institution + "  ")
    r1.bold = True
    p1.add_run(location + "\t")
    p1.add_run(period)
    p2 = doc.add_paragraph()
    _spacing(p2, before=0, after=30)
    p2.add_run(degree + "; ")
    p2.add_run(grade).bold = True


# ── Main builder ───────────────────────────────────────────

def build_docx(config: dict, job_role: str = "", job_description: str = "") -> dict:
    """
    Build a tailored DOCX resume.
    - Picks 2-3 best projects intelligently
    - Generates AI bullets via Groq
    Returns {"file_path": str, "resume_id": int, "filename": str}
    """
    if not DOCX_OK:
        return {"error": "python-docx not installed. Run: pip install python-docx"}

    profile    = config.get("profile", {})
    skills_cfg = config.get("skills", {})
    education  = config.get("education", [])
    certs_cfg  = config.get("certifications", [])

    gh       = db.get_latest_snapshot("github")
    cf       = db.get_latest_snapshot("codeforces")
    lc       = db.get_latest_snapshot("leetcode")
    all_proj = db.get_projects(limit=50, only_resume=True)   # fetch wide pool
    db_certs = db.get_certs()
    db_exp   = db.get_experience()

    all_certs = list(certs_cfg) + db_certs
    all_exp   = list(config.get("experience", [])) + db_exp

    # Smart selection: 2-3 best projects
    projs = _select_best_projects(all_proj, max_count=3)

    # Generate AI bullets for selected projects
    project_bullets = {}
    for proj in projs:
        print(f"[Resume] Generating bullets for: {proj['name']}")
        project_bullets[proj["name"]] = generate_resume_bullets(proj, config=config)

    # ── Build DOCX ─────────────────────────────────────────
    doc = Document()
    sec = doc.sections[0]
    sec.page_width    = Inches(8.27)
    sec.page_height   = Inches(11.70)
    sec.top_margin    = Inches(0.42)
    sec.bottom_margin = Inches(0.37)
    sec.left_margin   = Inches(0.43)
    sec.right_margin  = Inches(0.46)
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(10)

    # Name
    p_name = doc.add_paragraph()
    p_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _spacing(p_name, before=0, after=0)
    r = p_name.add_run(profile.get("name", ""))
    r.bold = True; r.font.size = Pt(18); r.font.color.rgb = NAME_COLOR

    # Contact row 1
    p1 = doc.add_paragraph()
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _spacing(p1, before=0, after=0); _right_tab(p1, 4.0)
    p1.add_run("LinkedIn: " + profile.get("linkedin", "") + "\t").font.size = Pt(9)
    p1.add_run("Email: " + profile.get("email", "")).font.size = Pt(9)

    # Contact row 2
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _spacing(p2, before=0, after=0); _right_tab(p2, 4.0)
    gh_url = gh.get("profile_url", f"github.com/{profile.get('github_username', '')}")
    p2.add_run("GitHub: " + gh_url + "\t").font.size = Pt(9)
    p2.add_run("Mobile: " + profile.get("phone", "")).font.size = Pt(9)

    # Skills
    _section(doc, "SKILLS")
    langs = list(skills_cfg.get("languages", []))
    for s in db.get_manual_skills():
        if s not in langs:
            langs.append(s)
    _skill_line(doc, "Languages",            ", ".join(langs))
    _skill_line(doc, "Web Technologies",     ", ".join(skills_cfg.get("frameworks", [])))
    _skill_line(doc, "Tools and Frameworks", ", ".join(skills_cfg.get("tools", [])))
    _skill_line(doc, "Core Competencies",    ", ".join(skills_cfg.get("cs_fundamentals", [])))
    if skills_cfg.get("soft_skills"):
        _skill_line(doc, "Soft Skills", ", ".join(skills_cfg["soft_skills"]))
    if job_role:
        _skill_line(doc, "Target Role", job_role)

    # Projects (smart-selected, AI bullets)
    _section(doc, "PROJECTS")
    if projs:
        for proj in projs:
            date_str = (proj.get("updated_at") or "")[:7]
            _project_header(doc, proj["name"] + " | GitHub", date_str)
            bullets = project_bullets.get(proj["name"], [])
            if not bullets:
                desc = proj.get("description", "")
                bullets = ([desc] if desc else []) + ["Tech: " + (proj.get("language") or "N/A")]
            for b in bullets:
                _bullet(doc, b)
    else:
        doc.add_paragraph().add_run(
            "Import GitHub projects from the GitHub Import page to fill this section."
        )

    # Training / Experience
    if all_exp:
        _section(doc, "TRAINING")
        for entry in all_exp:
            parts      = [x.strip() for x in entry.split("||")]
            title_line = parts[0]
            title_parts = [x.strip() for x in title_line.split("|")]
            p = doc.add_paragraph()
            _spacing(p, before=60, after=20); _right_tab(p, 7.3)
            if len(title_parts) >= 3:
                r1 = p.add_run("  " + title_parts[0] + "\t")
                r1.bold = True; r1.font.size = Pt(11); r1.font.color.rgb = LABEL_COLOR
                p.add_run(title_parts[2])
            else:
                r1 = p.add_run("  " + title_line)
                r1.bold = True; r1.font.color.rgb = LABEL_COLOR
            for b in parts[1:]:
                _bullet(doc, b)

    # Certifications
    if all_certs:
        _section(doc, "CERTIFICATES")
        for cert in all_certs:
            parts = [x.strip() for x in cert.split("|")]
            p = doc.add_paragraph()
            _spacing(p, before=0, after=20); _right_tab(p, 7.3)
            if len(parts) >= 3:
                p.add_run(parts[0] + " | " + parts[1] + "\t")
                p.add_run(parts[2])
            else:
                p.add_run(cert)

    # Achievements
    _section(doc, "ACHIEVEMENTS")
    if cf and cf.get("rating"):
        p = doc.add_paragraph()
        _spacing(p, before=60, after=0); _right_tab(p, 7.3)
        r1 = p.add_run("CodeForces: " + cf.get("profile_url", "") + "\t")
        r1.bold = True; r1.font.color.rgb = LABEL_COLOR
        _bullet(doc, f"Achieved rating of {cf.get('rating',0)} ({cf.get('rank','')}) through competitive programming.")
        _bullet(doc, f"Solved {cf.get('solved_count',0)}+ algorithmic problems across data structures, graphs, and DP.")

    if lc and lc.get("solved_total"):
        p = doc.add_paragraph()
        _spacing(p, before=60, after=0)
        r1 = p.add_run("LeetCode: " + lc.get("profile_url", ""))
        r1.bold = True; r1.font.color.rgb = LABEL_COLOR
        _bullet(doc, f"Solved {lc.get('solved_total',0)}+ problems: Easy {lc.get('solved_easy',0)}, Medium {lc.get('solved_medium',0)}, Hard {lc.get('solved_hard',0)}.")

    # Education
    _section(doc, "EDUCATION")
    for edu in education:
        grade = ("CGPA: " + str(edu["cgpa"])) if edu.get("cgpa") else edu.get("percentage", "")
        _edu_entry(doc,
            institution=edu.get("institution", ""),
            location=edu.get("location", ""),
            degree=edu.get("degree", ""),
            grade=grade,
            period=edu.get("year", ""))

    # Save
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    safe_name = (job_role or "General").replace(" ", "_").replace("/", "-")
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename  = f"{profile.get('name','Resume').replace(' ','_')}_{safe_name}_{timestamp}.docx"
    file_path = str(OUTPUT_DIR / filename)
    doc.save(file_path)

    resume_id = db.log_resume(
        title=filename,
        job_role=job_role or "General",
        file_path=file_path,
        ai_content=json.dumps(project_bullets),
    )

    print(f"[Resume] Saved -> {file_path}")
    return {"file_path": file_path, "resume_id": resume_id, "filename": filename}
