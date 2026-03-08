"""
modules/resume_engine.py
Builds the DOCX resume using AI-generated content from ai_writer.py.
Matches Gowtham's exact CV format: colors, margins, section order.
"""

import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).parent.parent))
import database as db

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    DOCX_OK = True
except ImportError:
    DOCX_OK = False
    print("[Resume] python-docx not installed. Run: pip install python-docx")

NAME_COLOR    = RGBColor(0x2F, 0x54, 0x96)
SECTION_COLOR = RGBColor(0x1F, 0x38, 0x64)
LABEL_COLOR   = RGBColor(0x00, 0x33, 0xCC)
LINK_COLOR    = RGBColor(0x00, 0x00, 0xFF)
BORDER_HEX    = "2F5496"


def _spacing(para, before=0, after=0):
    pPr = para._p.get_or_add_pPr()
    sp = OxmlElement('w:spacing')
    sp.set(qn('w:before'), str(before))
    sp.set(qn('w:after'), str(after))
    pPr.append(sp)


def _bottom_border(para):
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    b = OxmlElement('w:bottom')
    b.set(qn('w:val'), 'single')
    b.set(qn('w:sz'), '6')
    b.set(qn('w:space'), '1')
    b.set(qn('w:color'), BORDER_HEX)
    pBdr.append(b)
    pPr.append(pBdr)


def _right_tab(para, inches=7.3):
    pPr = para._p.get_or_add_pPr()
    tabs = OxmlElement('w:tabs')
    tab = OxmlElement('w:tab')
    tab.set(qn('w:val'), 'right')
    tab.set(qn('w:pos'), str(int(inches * 1440)))
    tabs.append(tab)
    pPr.append(tabs)


def _section(doc, title):
    p = doc.add_paragraph()
    _spacing(p, before=100, after=40)
    _bottom_border(p)
    r = p.add_run(title)
    r.font.size = Pt(11)
    r.font.color.rgb = SECTION_COLOR
    return p


def _bullet(doc, text):
    if not text or not text.strip():
        return
    p = doc.add_paragraph(style='List Paragraph')
    _spacing(p, before=0, after=20)
    p.add_run(text.strip())


def _skill_line(doc, label, value):
    if not value:
        return
    p = doc.add_paragraph()
    _spacing(p, before=0, after=20)
    r1 = p.add_run(label)
    r1.bold = True
    r1.font.color.rgb = LABEL_COLOR
    p.add_run(': ' + value)


def _project_header(doc, name, date):
    p = doc.add_paragraph()
    _spacing(p, before=80, after=20)
    _right_tab(p, 7.3)
    r1 = p.add_run('  ' + name + '\t')
    r1.bold = True
    r1.font.size = Pt(11)
    r1.font.color.rgb = LABEL_COLOR
    r2 = p.add_run(date)
    r2.font.size = Pt(10)


def _cert_line(doc, text, date=''):
    p = doc.add_paragraph()
    _spacing(p, before=0, after=0)
    _right_tab(p, 7.3)
    p.add_run(text + '\t')
    p.add_run(date)


def _platform_header(doc, platform, url, date_note=''):
    p = doc.add_paragraph()
    _spacing(p, before=60, after=0)
    _right_tab(p, 7.3)
    r1 = p.add_run(platform + ': ' + url + '\t')
    r1.bold = True
    r1.font.color.rgb = LINK_COLOR
    p.add_run(date_note)


def _edu_entry(doc, institution, location, degree, grade, period):
    p1 = doc.add_paragraph()
    _spacing(p1, before=40, after=0)
    _right_tab(p1, 7.3)
    r1 = p1.add_run(institution + '  ')
    r1.bold = True
    p1.add_run(location + '\t')
    p1.add_run(period)
    p2 = doc.add_paragraph()
    _spacing(p2, before=0, after=30)
    p2.add_run(degree + '; ')
    rb = p2.add_run(grade)
    rb.bold = True


# ══════════════════════════════════════════════════════════

def generate_resume(config: dict, output_path: str = None, force_ai=False) -> str:
    if not DOCX_OK:
        print("[Resume] python-docx not available.")
        return None

    # ── Pull AI-generated content ──────────────────────────
    from modules.ai_writer import generate_all_ai_content
    print("[Resume] Running AI writer...")
    ai = generate_all_ai_content(config, force_regenerate=force_ai)

    profile    = config.get("profile", {})
    skills_cfg = config.get("skills", {})
    education  = config.get("education", [])
    certs_cfg  = config.get("certifications", [])

    gh    = db.get_latest_snapshot("github")
    cf    = db.get_latest_snapshot("codeforces")
    lc    = db.get_latest_snapshot("leetcode")
    projs = db.get_projects(limit=6)
    db_certs  = db.get_certs()
    db_exp    = db.get_experience()

    all_certs = list(certs_cfg) + db_certs
    all_exp   = list(config.get("experience", [])) + db_exp

    if not output_path:
        output_path = config.get("resume", {}).get("output_path", "resume/Gowtham_Reddy_Resume.docx")
    if output_path.endswith(".pdf"):
        output_path = output_path.replace(".pdf", ".docx")
    Path(output_path).parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    sec = doc.sections[0]
    sec.page_width    = Inches(8.27)
    sec.page_height   = Inches(11.70)
    sec.top_margin    = Inches(0.42)
    sec.bottom_margin = Inches(0.37)
    sec.left_margin   = Inches(0.43)
    sec.right_margin  = Inches(0.46)

    doc.styles['Normal'].font.name = 'Calibri'
    doc.styles['Normal'].font.size = Pt(10)

    # ── NAME ──────────────────────────────────────────────
    p_name = doc.add_paragraph()
    p_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _spacing(p_name, before=0, after=0)
    r = p_name.add_run(profile.get("name", ""))
    r.bold = True
    r.font.size = Pt(18)
    r.font.color.rgb = NAME_COLOR

    # ── CONTACT ───────────────────────────────────────────
    p1 = doc.add_paragraph()
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _spacing(p1, before=0, after=0)
    _right_tab(p1, 4.0)
    r1 = p1.add_run('LinkedIn: ' + profile.get("linkedin", "") + '\t')
    r1.font.size = Pt(9)
    p1.add_run('Email: ' + profile.get("email", "")).font.size = Pt(9)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _spacing(p2, before=0, after=0)
    _right_tab(p2, 4.0)
    gh_url = gh.get("profile_url", "github.com/" + profile.get("github_username", ""))
    r3 = p2.add_run('GitHub: ' + gh_url + '\t')
    r3.font.size = Pt(9)
    p2.add_run('Mobile: ' + profile.get("phone", "")).font.size = Pt(9)

    # ══ SKILLS ════════════════════════════════════════════
    _section(doc, 'SKILLS')
    langs = list(skills_cfg.get("languages", []))
    for s in db.get_manual_skills():
        if s not in langs:
            langs.append(s)
    _skill_line(doc, 'Languages',            ', '.join(langs))
    _skill_line(doc, 'Web Technologies',     ', '.join(skills_cfg.get("frameworks", [])))
    _skill_line(doc, 'Tools and Frameworks', ', '.join(skills_cfg.get("tools", [])))
    _skill_line(doc, 'Core Competencies',    ', '.join(skills_cfg.get("cs_fundamentals", [])))
    if skills_cfg.get("soft_skills"):
        _skill_line(doc, 'Soft Skills', ', '.join(skills_cfg["soft_skills"]))

    # ══ PROJECTS (AI-written bullets) ═════════════════════
    _section(doc, 'PROJECTS')
    if projs:
        for proj in projs:
            date_str = proj.get("updated_at", "")[:7] if proj.get("updated_at") else ""
            _project_header(doc, proj["name"] + " | GitHub", date_str)

            # Get AI bullets for this project
            bullets = ai["projects"].get(proj["name"], [])
            if not bullets:
                desc = proj.get("description", "")
                if desc:
                    bullets = [desc, "Tech: " + (proj.get("language") or "N/A")]
            for b in bullets:
                _bullet(doc, b)
    else:
        doc.add_paragraph().add_run("Run /updatestats to pull your GitHub projects.")

    # ══ TRAINING (AI-enhanced) ════════════════════════════
    if all_exp:
        _section(doc, 'TRAINING')
        for entry in all_exp:
            # First line: title + date (split by first ||)
            parts = [x.strip() for x in entry.split("||")]
            title_line = parts[0]
            title_parts = [x.strip() for x in title_line.split("|")]

            p = doc.add_paragraph()
            _spacing(p, before=60, after=20)
            _right_tab(p, 7.3)
            if len(title_parts) >= 3:
                r1 = p.add_run('  ' + title_parts[0] + '\t')
                r1.bold = True
                r1.font.size = Pt(11)
                r1.font.color.rgb = LABEL_COLOR
                p.add_run(title_parts[2])
            else:
                r1 = p.add_run('  ' + title_line)
                r1.bold = True
                r1.font.color.rgb = LABEL_COLOR

            # AI-enhanced bullets
            ai_bullets = ai["training"].get(entry, [])
            if ai_bullets:
                for b in ai_bullets:
                    _bullet(doc, b)
            else:
                # Fallback to raw bullets after ||
                for b in parts[1:]:
                    _bullet(doc, b)

    # ══ CERTIFICATES ══════════════════════════════════════
    if all_certs:
        _section(doc, 'CERTIFICATES')
        for cert in all_certs:
            parts = [x.strip() for x in cert.split('|')]
            if len(parts) >= 3:
                _cert_line(doc, parts[0] + ' | ' + parts[1], parts[2])
            elif len(parts) == 2:
                _cert_line(doc, parts[0] + ' | ' + parts[1])
            else:
                _cert_line(doc, cert)

    # ══ ACHIEVEMENTS (AI-written, auto-updated) ════════════
    _section(doc, 'ACHIEVEMENTS')
    ach = ai.get("achievements", {})

    if cf:
        _platform_header(doc, 'CodeForces',
            cf.get("profile_url", ""), 'Since January 2026')
        for b in ach.get("codeforces", []):
            _bullet(doc, b)

    if lc:
        _platform_header(doc, 'LeetCode',
            lc.get("profile_url", ""), 'Since October 2025')
        for b in ach.get("leetcode", []):
            _bullet(doc, b)

    if gh:
        _platform_header(doc, 'GitHub', gh.get("profile_url", ""))
        for b in ach.get("github", []):
            _bullet(doc, b)

    # ══ EDUCATION ═════════════════════════════════════════
    _section(doc, 'EDUCATION')
    for edu in education:
        grade = ("CGPA: " + str(edu["cgpa"])) if edu.get("cgpa") else edu.get("percentage", "")
        _edu_entry(doc,
            institution=edu.get("institution", ""),
            location=edu.get("location", ""),
            degree=edu.get("degree", ""),
            grade=grade,
            period=edu.get("year", ""))

    doc.save(output_path)
    db.log_resume(output_path)
    print(f"[Resume] ✓ Saved → {output_path}")
    return output_path
